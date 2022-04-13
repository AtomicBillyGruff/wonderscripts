from docx import Document
import os


def main():
    template_file_path = '/home/kali/Downloads/coverletter.docx'
    output_file_path = '/home/kali/Downloads/Jacob-Simpson-CoverLetter.docx'

    variables = {

        "${COMPANY}": "Awesome Company",
        "${LOCATION}": "828 Bolo Road",
        "${CITY}": "Bethesda",
        "${STATE}": "Maryland",
        "${ZIP}": "643534",
        "${TITLE}": "Cybersecurity Artist",
        "${EVENT}": "meeting ...",
        "${ADJECTIVE}": "thrilled",
        "${PHILOS}": "Network Security",
        "${PHILOS2}": "Network Security",
        "${PHILOS3}": "Network Security",

        "${EXPERIENCE1}": "Washington University Cybersecurity Bootcamp",
        "${EXPERIENCE2}": "Community Multimedia Specialist",
        "${SKILLS}": "in Network Security, Application Vulnerability Testing, and Metric Logging",
        "${ADDRESS}": "jacb.simp@gmail.com",
        "${PHONE}": "314-677-7719",
        "${EMAIL}": "jacb.simp@gmail.com",
        "${NAME}": "Jacob Simpson",
        "${DATE}": "04/02/2022",
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == '__main__':
    main()